"""
Test that the parser extracts tables (not just paragraphs).

RFP requirements and the KB's project tables live in tables, which the old
paragraph-only parser dropped. Builds a small DOCX with a table on the fly so the
test is self-contained and reproducible.

Usage:
    python notebooks/test_parser_tables.py
"""
import shutil
import tempfile
from pathlib import Path

import docx

from src.ingestion.parser import parse_file, _rows_to_markdown


def main() -> None:
    # Unit check on the Markdown table renderer.
    md = _rows_to_markdown([["Requirement", "Priority"], ["99.9% uptime", "High"]])
    assert "| Requirement | Priority |" in md
    assert "| --- | --- |" in md
    assert "99.9% uptime" in md
    print("rows_to_markdown: OK")

    # Build a DOCX with a paragraph + a requirements table.
    tmp = Path(tempfile.mkdtemp(prefix="parsetest_"))
    path = tmp / "sample.docx"
    d = docx.Document()
    d.add_paragraph("Statement of requirements")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Requirement"
    table.rows[0].cells[1].text = "Priority"
    table.rows[1].cells[0].text = "High availability SQL database"
    table.rows[1].cells[1].text = "Mandatory"
    d.save(str(path))

    try:
        docs = parse_file(path, source_type="incoming_rfp")
        text = docs[0].page_content
        assert "Statement of requirements" in text, "paragraph text missing"
        assert "High availability SQL database" in text, "table cell missing"
        assert "|" in text, "table not rendered as Markdown"
        assert docs[0].metadata["source_type"] == "incoming_rfp"
        print("docx table extraction: OK")
        print(f"  extracted {len(text)} chars including the requirements table")
        print("\nTest completed successfully.")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    main()
